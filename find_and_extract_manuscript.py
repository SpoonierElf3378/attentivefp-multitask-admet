"""
Find the manuscript .docx and pull the sections needed for review items
1 (abstract), 2 (Table 1 caption/Delta column), 3 (reference list + citation
order), 10 (duplicate sentence, long sentences, Table 4 row order).

Run: pip install python-docx --no-deps  (if not already installed)
Then: python find_and_extract_manuscript.py
"""
import glob
import os

try:
    from docx import Document
except ImportError:
    print("Need python-docx. Run: pip install python-docx")
    raise SystemExit(1)

candidates = glob.glob(os.path.join("**", "*.docx"), recursive=True)
print(f"Found {len(candidates)} .docx files:")
for c in candidates:
    print(" ", c)

if not candidates:
    print("\nNo .docx found. If your manuscript is elsewhere, edit MANUSCRIPT_PATH below and rerun.")
    raise SystemExit(0)

# Heuristic: pick the biggest one (most likely the full manuscript, not a cover letter)
MANUSCRIPT_PATH = max(candidates, key=os.path.getsize)
print(f"\nUsing largest file as manuscript: {MANUSCRIPT_PATH}\n")

doc = Document(MANUSCRIPT_PATH)

print("=" * 80)
print("PARAGRAPH COUNT:", len(doc.paragraphs))
print("TABLE COUNT:", len(doc.tables))
print("=" * 80)

print("\n--- First 5 paragraphs (likely title/abstract area) ---")
for i, p in enumerate(doc.paragraphs[:15]):
    if p.text.strip():
        print(f"[{i}] {p.text[:200]}")

print("\n--- Searching all paragraphs for 'Abstract' ---")
for i, p in enumerate(doc.paragraphs):
    if "abstract" in p.text.lower():
        print(f"[{i}] {p.text[:300]}")
        # print next few paragraphs (likely abstract body)
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f"    [{j}] {doc.paragraphs[j].text}")

print("\n--- Table dump (index, first row = likely headers) ---")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:3]):
        cells = [c.text.strip() for c in row.cells]
        print(f"  row{ri}: {cells}")

print("\n--- References section (searching for 'References') ---")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower() == "references":
        print(f"References header at paragraph {i}")
        for j in range(i + 1, min(i + 25, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f"  [{j}] {doc.paragraphs[j].text}")
        break
